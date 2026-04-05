"""
Convert VitalSignLogger_Firmware_Textbook.md to a print-ready .docx file.
Run: python Python/textbook_to_docx.py
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH   = "VitalSignLogger_Firmware_Textbook.md"
DOCX_PATH = "VitalSignLogger_Firmware_Textbook.docx"

# ── Colour palette ────────────────────────────────────────────────────────────
C_HEADING1  = RGBColor(0x1F, 0x49, 0x7D)   # dark blue
C_HEADING2  = RGBColor(0x2E, 0x74, 0xB5)   # mid blue
C_HEADING3  = RGBColor(0x37, 0x85, 0xC4)   # lighter blue
C_CODE_BG   = "F2F2F2"                      # light grey hex string
C_CODE_TXT  = RGBColor(0x17, 0x17, 0x17)
C_TH_BG     = RGBColor(0x1F, 0x49, 0x7D)   # table header bg (dark blue)
C_TH_TXT    = RGBColor(0xFF, 0xFF, 0xFF)   # table header text (white)
C_RULE      = RGBColor(0x1F, 0x49, 0x7D)
C_TITLE     = RGBColor(0x1F, 0x49, 0x7D)
C_SUBTITLE  = RGBColor(0x59, 0x59, 0x59)


def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "{:02X}{:02X}{:02X}".format(rgb[0], rgb[1], rgb[2]))
    tcPr.append(shd)


def add_horizontal_rule(doc):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pb   = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(*C_RULE))
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_after = Pt(6)


def set_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.5):
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


def style_heading(para, level: int):
    colours = {1: C_HEADING1, 2: C_HEADING2, 3: C_HEADING3}
    sizes   = {1: 18, 2: 14, 3: 12}
    for run in para.runs:
        run.bold           = True
        run.font.color.rgb = colours.get(level, C_HEADING2)
        run.font.size      = Pt(sizes.get(level, 12))
    para.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    para.paragraph_format.space_after  = Pt(4)


def apply_inline(para, text: str):
    """Write text with **bold** and `code` inline markup rendered correctly."""
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts   = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = para.add_run(part[1:-1])
            r.font.name      = "Courier New"
            r.font.size      = Pt(9)
            r.font.color.rgb = C_CODE_TXT
        else:
            para.add_run(part)


def add_code_block(doc, lines):
    """Add a shaded monospace block for fenced code."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  C_CODE_BG)
        pPr.append(shd)
        r = p.add_run(line if line else " ")
        r.font.name      = "Courier New"
        r.font.size      = Pt(8.5)
        r.font.color.rgb = C_CODE_TXT
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def parse_table(lines):
    headers = []
    rows    = []
    for raw in lines:
        raw = raw.strip()
        if not raw.startswith("|"):
            continue
        cells = [c.strip() for c in raw.strip("|").split("|")]
        if all(re.fullmatch(r":?-+:?", c.strip()) for c in cells):
            continue
        if not headers:
            headers = cells
        else:
            rows.append(cells)
    return headers, rows


def add_table(doc, headers, rows):
    ncols = len(headers)
    tbl   = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style = "Table Grid"

    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold           = True
            run.font.color.rgb = C_TH_TXT
            run.font.size      = Pt(9)
        set_cell_bg(hdr_cells[i], C_TH_BG)

    for r_idx, row in enumerate(rows):
        cells = tbl.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row):
            if c_idx < ncols:
                p = cells[c_idx].paragraphs[0]
                for run in p.runs:
                    run.text = ""
                apply_inline(p, cell_text)
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_title_page(doc):
    """Create a professional title page."""
    # Top spacing
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("VitalSignLogger")
    r.bold           = True
    r.font.size      = Pt(36)
    r.font.color.rgb = C_TITLE
    r.font.name      = "Calibri"

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Firmware Textbook")
    r.bold           = True
    r.font.size      = Pt(24)
    r.font.color.rgb = C_SUBTITLE
    r.font.name      = "Calibri"

    doc.add_paragraph()

    # Tagline
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A line-by-line guide to every embedded C concept in this project")
    r.italic         = True
    r.font.size      = Pt(13)
    r.font.color.rgb = C_SUBTITLE
    r.font.name      = "Calibri"

    # Spacing
    for _ in range(6):
        doc.add_paragraph()

    # Details
    details = [
        ("Platform", "STM32U575ZIT6Q  (Cortex-M33, 160 MHz)"),
        ("Board", "NUCLEO-U575ZI-Q"),
        ("IDE", "STM32CubeIDE 1.19.0"),
        ("RTOS", "FreeRTOS (CMSIS-RTOS2)"),
        ("ML", "STM32Cube.AI  (3-class motion classifier)"),
    ]

    tbl = doc.add_table(rows=len(details), cols=2)
    tbl.style = "Table Grid"
    # Remove borders for a clean look
    for row_idx, (label, value) in enumerate(details):
        tbl.rows[row_idx].cells[0].text = label
        tbl.rows[row_idx].cells[1].text = value
        for ci in range(2):
            cell = tbl.rows[row_idx].cells[ci]
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(11)
                run.font.name = "Calibri"
            if ci == 0:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = C_HEADING1
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            # light or no background
            if row_idx % 2 == 0:
                set_cell_bg(cell, RGBColor(0xF2, 0xF2, 0xF2))

    doc.add_paragraph()

    # Page break after title
    add_page_break(doc)


def add_page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def main():
    with open(MD_PATH, encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    set_margins(doc)

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)

    # Add title page
    add_title_page(doc)

    i = 0
    n = len(lines)

    # Skip the markdown title lines (# VitalSignLogger..., **A line-by-line...**)
    # since we already have a title page
    while i < n:
        raw = lines[i].rstrip("\n").strip()
        if raw.startswith("# ") or raw.startswith("**A line-by-line"):
            i += 1
            continue
        if raw == "---":
            i += 1
            continue
        if raw == "":
            i += 1
            continue
        break  # first real content line

    # Track chapter numbers for page breaks (break before each ## heading)
    chapter_count = 0

    while i < n:
        raw = lines[i].rstrip("\n")

        # ── Blank line ──
        if raw.strip() == "":
            i += 1
            continue

        # ── Horizontal rule ---  ──
        if re.fullmatch(r"-{3,}", raw.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Fenced code block ``` ──
        if raw.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1  # skip closing ```
            add_code_block(doc, code_lines)
            continue

        # ── Table ──
        if raw.strip().startswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip("\n"))
                i += 1
            headers, rows = parse_table(table_lines)
            if headers:
                add_table(doc, headers, rows)
            continue

        # ── Headings # ## ### ──
        m = re.match(r"^(#{1,3})\s+(.*)", raw)
        if m:
            level = len(m.group(1))
            text  = m.group(2).strip()
            # Strip markdown link syntax [text](#anchor) → text
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            # Strip trailing bold markers
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)

            # Page break before each chapter (## heading), except first
            if level == 2:
                chapter_count += 1
                if chapter_count > 1:
                    add_page_break(doc)

            hmap = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
            para = doc.add_paragraph(style=hmap.get(level, "Heading 2"))
            para.clear()
            para.add_run(text)
            style_heading(para, level)
            i += 1
            continue

        # ── Blockquote > ──
        if raw.strip().startswith(">"):
            text = re.sub(r"^>\s*", "", raw.strip())
            para = doc.add_paragraph(style="Quote")
            para.clear()
            apply_inline(para, text)
            for run in para.runs:
                run.font.size   = Pt(9.5)
                run.font.italic = True
            i += 1
            continue

        # ── Bullet list ──
        if re.match(r"^(\s*)[-*+]\s+", raw):
            indent = len(re.match(r"^(\s*)", raw).group(1))
            text   = re.sub(r"^\s*[-*+]\s+", "", raw)
            level_style = "List Bullet 2" if indent >= 2 else "List Bullet"
            para = doc.add_paragraph(style=level_style)
            para.clear()
            apply_inline(para, text)
            para.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # ── Numbered list ──
        if re.match(r"^\d+\.\s+", raw):
            text = re.sub(r"^\d+\.\s+", "", raw)
            para = doc.add_paragraph(style="List Number")
            para.clear()
            apply_inline(para, text)
            para.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # ── Normal paragraph ──
        para = doc.add_paragraph()
        apply_inline(para, raw.strip())
        para.paragraph_format.space_after = Pt(4)
        i += 1

    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}  ({chapter_count} chapters)")


if __name__ == "__main__":
    main()
