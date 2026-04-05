"""
Convert ML_PIPELINE_DEEP_DIVE.md to a print-ready .docx file.
Run: python Python/md_to_docx.py
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

MD_PATH   = "Python/ML_PIPELINE_DEEP_DIVE.md"
DOCX_PATH = "Python/ML_PIPELINE_DEEP_DIVE.docx"

# ── Colour palette ────────────────────────────────────────────────────────────
C_HEADING1  = RGBColor(0x1F, 0x49, 0x7D)   # dark blue
C_HEADING2  = RGBColor(0x2E, 0x74, 0xB5)   # mid blue
C_HEADING3  = RGBColor(0x2E, 0x74, 0xB5)
C_CODE_BG   = RGBColor(0xF2, 0xF2, 0xF2)   # light grey
C_CODE_TXT  = RGBColor(0x17, 0x17, 0x17)
C_TH_BG     = RGBColor(0x1F, 0x49, 0x7D)   # table header bg (dark blue)
C_TH_TXT    = RGBColor(0xFF, 0xFF, 0xFF)   # table header text (white)
C_RULE      = RGBColor(0x1F, 0x49, 0x7D)


def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "{:02X}{:02X}{:02X}".format(rgb[0], rgb[1], rgb[2]))
    tcPr.append(shd)


def add_horizontal_rule(doc):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pb   = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(*C_RULE))
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_after = Pt(6)


def set_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.5):
    """Set page margins in centimetres."""
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


def style_heading(para, level: int):
    colours = {1: C_HEADING1, 2: C_HEADING2, 3: C_HEADING3}
    sizes   = {1: 18, 2: 14, 3: 12}
    run = para.runs[0] if para.runs else para.add_run(para.text)
    run.bold       = True
    run.font.color.rgb = colours.get(level, C_HEADING2)
    run.font.size  = Pt(sizes.get(level, 12))
    para.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    para.paragraph_format.space_after  = Pt(4)


def apply_inline(para, text: str):
    """
    Write text with **bold** and `code` inline markup rendered correctly.
    Each segment is added as a separate Run so formatting is applied.
    """
    # Split on **bold** and `code` tokens
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts   = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = para.add_run(part[1:-1])
            r.font.name  = "Courier New"
            r.font.size  = Pt(9)
            r.font.color.rgb = C_CODE_TXT
        else:
            para.add_run(part)


def add_code_block(doc, lines):
    """Add a shaded monospace block for fenced code."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        # Grey shading via paragraph shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F2F2F2")
        pPr.append(shd)
        r = p.add_run(line if line else " ")
        r.font.name  = "Courier New"
        r.font.size  = Pt(8.5)
        r.font.color.rgb = C_CODE_TXT
    # small gap after block
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def parse_table(lines):
    """
    Parse a markdown table (list of raw lines that form the table).
    Returns (headers: list[str], rows: list[list[str]]).
    The separator line (---) is skipped.
    """
    headers = []
    rows    = []
    for raw in lines:
        raw = raw.strip()
        if not raw.startswith("|"):
            continue
        cells = [c.strip() for c in raw.strip("|").split("|")]
        if all(re.fullmatch(r"-+", c.strip("-")) for c in cells):
            continue   # separator row
        if not headers:
            headers = cells
        else:
            rows.append(cells)
    return headers, rows


def add_table(doc, headers, rows):
    ncols = len(headers)
    tbl   = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style = "Table Grid"

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold            = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb  = C_TH_TXT
        hdr_cells[i].paragraphs[0].runs[0].font.size       = Pt(9)
        set_cell_bg(hdr_cells[i], C_TH_BG)

    # Data rows
    for r_idx, row in enumerate(rows):
        cells = tbl.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row):
            if c_idx < ncols:
                p = cells[c_idx].paragraphs[0]
                # Clear default empty run
                for run in p.runs:
                    run.text = ""
                apply_inline(p, cell_text)
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def main():
    with open(MD_PATH, encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    set_margins(doc)

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    i = 0
    n = len(lines)

    while i < n:
        raw = lines[i].rstrip("\n")

        # ── Blank line ───────────────────────────────────────────────────────
        if raw.strip() == "":
            i += 1
            continue

        # ── Horizontal rule ---  ──────────────────────────────────────────────
        if re.fullmatch(r"-{3,}", raw.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Fenced code block ```  ────────────────────────────────────────────
        if raw.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1  # skip closing ```
            add_code_block(doc, code_lines)
            continue

        # ── Table  ────────────────────────────────────────────────────────────
        if raw.strip().startswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip("\n"))
                i += 1
            headers, rows = parse_table(table_lines)
            if headers:
                add_table(doc, headers, rows)
            continue

        # ── Headings # ## ### ─────────────────────────────────────────────────
        m = re.match(r"^(#{1,3})\s+(.*)", raw)
        if m:
            level = len(m.group(1))
            text  = m.group(2).strip()
            # Strip trailing bold/italic markers
            text  = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
            hmap  = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
            para  = doc.add_paragraph(style=hmap.get(level, "Heading 2"))
            para.clear()
            run   = para.add_run(text)
            style_heading(para, level)
            i += 1
            continue

        # ── Blockquote >  ─────────────────────────────────────────────────────
        if raw.strip().startswith(">"):
            text = re.sub(r"^>\s*", "", raw.strip())
            para = doc.add_paragraph(style="Quote")
            para.clear()
            apply_inline(para, text)
            for run in para.runs:
                run.font.size  = Pt(9.5)
                run.font.italic = True
            i += 1
            continue

        # ── Bullet list  ──────────────────────────────────────────────────────
        if re.match(r"^(\s*)[-*+]\s+", raw):
            indent = len(re.match(r"^(\s*)", raw).group(1))
            text   = re.sub(r"^\s*[-*+]\s+", "", raw)
            level_style = "List Bullet 2" if indent >= 2 else "List Bullet"
            para   = doc.add_paragraph(style=level_style)
            para.clear()
            apply_inline(para, text)
            para.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # ── Numbered list  ────────────────────────────────────────────────────
        if re.match(r"^\d+\.\s+", raw):
            text = re.sub(r"^\d+\.\s+", "", raw)
            para = doc.add_paragraph(style="List Number")
            para.clear()
            apply_inline(para, text)
            para.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # ── Normal paragraph  ─────────────────────────────────────────────────
        para = doc.add_paragraph()
        apply_inline(para, raw.strip())
        para.paragraph_format.space_after = Pt(4)
        i += 1

    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")


if __name__ == "__main__":
    main()
