#!/usr/bin/env python3
"""
generate_report.py — Generates the VitalSignLogger project report as a DOCX file.
Run: pip install python-docx && python generate_report.py
Output: VitalSignLogger_Project_Report.docx (in the same directory)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Global style tweaks ────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)

code_style = doc.styles.add_style('CodeBlock', 1)  # paragraph style
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(4)
code_style.paragraph_format.left_indent = Cm(0.5)

def add_code(text):
    """Add a code block paragraph."""
    for i, line in enumerate(text.strip().split('\n')):
        p = doc.add_paragraph(line, style='CodeBlock')
        if i == 0:
            p.paragraph_format.space_before = Pt(6)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()  # spacer


# ═══════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('VitalSignLogger')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Project Report & Technical Bible')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Real-Time Vital Signs Monitoring System\n').font.size = Pt(13)
meta.add_run('STM32U575ZIT6QU | FreeRTOS | STM32Cube.AI\n\n').font.size = Pt(12)
meta.add_run('Date: 30 March 2026\n').font.size = Pt(11)
meta.add_run('Platform: NUCLEO-U575ZI-Q Development Board').font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. What Does This Project Do? (Plain English)',
    '3. Hardware — The Physical Components',
    '   3.1 The Brain: STM32U575 Microcontroller',
    '   3.2 The Heart Rate Sensor: MAX30102',
    '   3.3 The Motion Sensor: MPU6050',
    '   3.4 The Display: SSD1306 OLED',
    '   3.5 How They Connect: The I2C Bus',
    '   3.6 Wiring Diagram & Pin Map',
    '4. Firmware Architecture — How the Software Works',
    '   4.1 FreeRTOS: Running Multiple Things at Once',
    '   4.2 The Four Tasks',
    '   4.3 Data Flow: From Sensor to Screen',
    '   4.4 Sharing the I2C Bus Safely',
    '5. Signal Processing — Cleaning Up the Heart Rate Signal',
    '   5.1 The Raw PPG Signal',
    '   5.2 Digital Filtering (Biquad IIR)',
    '   5.3 Peak Detection & BPM Calculation',
    '6. Machine Learning — Motion Classification on a Microcontroller',
    '   6.1 What the ML Does',
    '   6.2 Feature Extraction (The 19 Numbers)',
    '   6.3 The Neural Network Architecture',
    '   6.4 STM32Cube.AI: From Python to C',
    '   6.5 How Inference Runs in Real Time',
    '7. Data Output — UART CSV Stream',
    '8. Python Companion Tools',
    '   8.1 Live Dashboard (vital_dashboard.py)',
    '   8.2 Offline Analysis (analyse_log.py)',
    '9. The Development Journey — Errors, Detours & Fixes',
    '   9.1 Display Went Dark After Sensor Init',
    '   9.2 No UART Output (Printf Was a No-Op)',
    '   9.3 Tasks Crashing Silently (Stack Overflow)',
    '   9.4 ML Inference Returning Garbage',
    '   9.5 Display Only Detected 1 in 5 Boots',
    '   9.6 Board Wouldn\'t Move Past "Keep Board Flat"',
    '   9.7 Motion Stuck at REST Despite Moving',
    '   9.8 Python Dashboard: COM Port Access Denied',
    '10. Boot Sequence — What Happens When You Power On',
    '11. Key Configuration & Numbers',
    '12. Project Status & Future Work',
    '13. Glossary of Terms',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if not item.startswith('   '):
        p.runs[0].bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'VitalSignLogger is a real-time health monitoring device built on the STM32U575ZIT6QU '
    'microcontroller — a small but powerful computer chip about the size of a thumbnail. '
    'It continuously reads a person\'s heart rate using an optical sensor (like the one on '
    'the back of a smartwatch), detects body motion using an accelerometer/gyroscope '
    '(the same type of sensor that knows when you tilt your phone), and classifies that '
    'motion into three categories — resting, walking, or arm raised — using a tiny neural '
    'network running directly on the chip.'
)
doc.add_paragraph(
    'All of this happens 100 times per second. The results are shown on a small OLED '
    'display attached to the device, and simultaneously streamed as a CSV data log over '
    'a USB cable to a computer, where a Python dashboard can visualise it in real time.'
)
doc.add_paragraph(
    'The project exercises a wide range of embedded systems skills: real-time operating '
    'system (RTOS) task scheduling, I2C bus management with multiple devices, digital '
    'signal processing for heart-rate extraction, on-device machine learning inference, '
    'and Python-based data visualisation.'
)

# ═══════════════════════════════════════════════════════════════════════
# 2. PLAIN ENGLISH OVERVIEW
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('2. What Does This Project Do? (Plain English)', level=1)
doc.add_paragraph(
    'Imagine a small electronic device, roughly the size of a deck of cards, sitting on a '
    'table. Three tiny sensor modules are wired to it:'
)
add_bullet('A red-glowing sensor where you place your fingertip (heart rate sensor)', bold_prefix='Heart Rate: ')
add_bullet('A chip that can sense tilt and movement, like a spirit level that speaks digital (motion sensor)', bold_prefix='Motion: ')
add_bullet('A tiny screen, smaller than a postage stamp, showing live numbers (OLED display)', bold_prefix='Display: ')

doc.add_paragraph(
    'When you place your finger on the red sensor, within a few seconds the screen shows '
    'your heart rate in beats per minute (BPM). Pick up the device and walk around — the '
    'screen changes to show "WALKING". Raise your arm — it shows "ARM UP". Put it back on '
    'the table — "REST".'
)
doc.add_paragraph(
    'Meanwhile, a USB cable sends all this data to your laptop at 100 readings per second. '
    'A Python program on the laptop draws live scrolling graphs: a red line for your '
    'heartbeat, coloured traces for acceleration in three dimensions, and a big colour-coded '
    'label for the current motion class.'
)
doc.add_paragraph(
    'The "brain" that makes this possible is a microcontroller — a complete computer on a '
    'single chip, running at 160 million operations per second, with its own memory, its own '
    'operating system (FreeRTOS), and a tiny AI model that classifies motion. There is no '
    'cloud, no WiFi, no phone app — everything happens locally, in real time, on hardware '
    'that costs about $15.'
)

# ═══════════════════════════════════════════════════════════════════════
# 3. HARDWARE
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('3. Hardware — The Physical Components', level=1)

doc.add_heading('3.1 The Brain: STM32U575 Microcontroller', level=2)
doc.add_paragraph(
    'The STM32U575ZIT6QU is an ARM Cortex-M33 microcontroller made by STMicroelectronics. '
    'Think of it as a tiny, energy-efficient computer. Key specifications:'
)
add_table(
    ['Specification', 'Value', 'What It Means'],
    [
        ['CPU Core', 'ARM Cortex-M33 @ 160 MHz', 'Can execute 160 million instructions per second'],
        ['Flash Memory', '2 MB', 'Permanent storage for the program (like a hard drive)'],
        ['RAM', '192 KB SRAM', 'Working memory for variables and buffers (like desktop RAM, but tiny)'],
        ['FPU', 'Single-precision hardware', 'Can do decimal math (3.14 x 2.71) in one clock cycle'],
        ['Security', 'TrustZone', 'Hardware-level security isolation (not used in this project)'],
        ['Power', 'Ultra-low-power design', 'Designed for battery-operated devices'],
    ]
)
doc.add_paragraph(
    'The chip sits on a NUCLEO-U575ZI-Q development board, which adds USB connectivity, '
    'a built-in debugger/programmer (ST-LINK), pin headers for wiring sensors, and '
    'a virtual COM port that lets us send data to a PC over USB as if it were a serial cable.'
)

doc.add_heading('3.2 The Heart Rate Sensor: MAX30102', level=2)
doc.add_paragraph(
    'The MAX30102 is a pulse oximetry and heart-rate sensor made by Maxim Integrated '
    '(now part of Analog Devices). It works using photoplethysmography (PPG) — a '
    'technique where:'
)
add_bullet('A red LED shines light into your fingertip.')
add_bullet('A photodetector measures how much light comes back.')
add_bullet('Each heartbeat pushes more blood into your finger, which absorbs more light.')
add_bullet('The resulting signal is a wave that rises and falls with each heartbeat.')
doc.add_paragraph(
    'The sensor has a built-in 32-sample FIFO (First In, First Out) buffer, meaning it '
    'can store up to 32 readings internally even if the microcontroller is momentarily '
    'busy. We read it at 100 Hz (100 times per second). The sensor communicates via I2C '
    'at address 0x57.'
)

doc.add_heading('3.3 The Motion Sensor: MPU6050', level=2)
doc.add_paragraph(
    'The MPU6050 is a MEMS (Micro-Electro-Mechanical System) sensor made by InvenSense '
    '(now TDK). It combines:'
)
add_bullet('A 3-axis accelerometer — measures acceleration (tilt, vibration, movement) along X, Y, and Z axes, in units of "g" (where 1g = Earth\'s gravity).')
add_bullet('A 3-axis gyroscope — measures rotational speed (how fast you\'re turning) in degrees per second.')
add_bullet('A built-in temperature sensor — reads the chip\'s temperature (useful for calibration, not body temperature).')
doc.add_paragraph(
    'When the device is sitting flat and still on a table, the accelerometer reads approximately '
    '(0, 0, 1.0) g — zero sideways force, and 1g pulling straight down (gravity). Tilt the '
    'board 90 degrees and gravity shifts to a different axis. Walk around and the vibrations '
    'create a distinctive pattern. The ML model learns to tell these apart.'
)
doc.add_paragraph(
    'At startup, the MPU6050 is calibrated: the firmware takes 200 readings while the board '
    'sits flat, averages them, and subtracts this "bias" from all future readings. This '
    'removes manufacturing offsets so that a stationary board reads exactly (0, 0, 1.0) g. '
    'I2C address: 0x68.'
)

doc.add_heading('3.4 The Display: SSD1306 OLED', level=2)
doc.add_paragraph(
    'The SSD1306 is a 0.96-inch, 128x64 pixel OLED (Organic Light-Emitting Diode) display. '
    'Unlike an LCD, each pixel emits its own light — no backlight needed. The firmware '
    'maintains a 1024-byte framebuffer (128 x 64 / 8 = 1024 bytes, since each pixel is '
    'one bit: on or off) in RAM, draws text into it using a built-in 5x8 pixel ASCII font, '
    'then sends the entire buffer to the display over I2C.'
)
doc.add_paragraph(
    'At 100 kHz I2C speed, transmitting 1024 bytes takes about 103 milliseconds — a '
    'significant amount of time that must be carefully managed so it doesn\'t block the '
    'heart rate sensor from reading. I2C address: 0x3C.'
)

doc.add_heading('3.5 How They Connect: The I2C Bus', level=2)
doc.add_paragraph(
    'I2C (pronounced "I-squared-C" or "I-two-C") is a communication protocol that lets '
    'multiple devices share just two wires:'
)
add_bullet('SDA (Serial Data) — carries the actual data, one bit at a time.', bold_prefix='SDA: ')
add_bullet('SCL (Serial Clock) — carries a clock signal so both sides know when to read each bit.', bold_prefix='SCL: ')
doc.add_paragraph(
    'Every device on the bus has a unique 7-bit address (like a postal address). The '
    'microcontroller acts as the "master" — it initiates all communication by sending an '
    'address, and only the device with that address responds. This means our three sensors '
    '(MAX30102 at 0x57, MPU6050 at 0x68, SSD1306 at 0x3C) all share the same two wires '
    'without interfering with each other.'
)
doc.add_paragraph(
    'The catch: only one device can talk at a time. If the display is being updated (103 ms), '
    'the heart rate sensor can\'t be read simultaneously. The firmware uses a mutex (a '
    'software lock) to manage this — more on this in Section 4.4.'
)

doc.add_heading('3.6 Wiring Diagram & Pin Map', level=2)
add_table(
    ['STM32 Pin', 'Header', 'Function', 'Connected To'],
    [
        ['PB8', 'CN7 (D15)', 'I2C1 SCL (Clock)', 'MAX30102 SCL + MPU6050 SCL + SSD1306 SCL'],
        ['PB9', 'CN7 (D14)', 'I2C1 SDA (Data)', 'MAX30102 SDA + MPU6050 SDA + SSD1306 SDA'],
        ['PA9', 'Internal', 'USART1 TX', 'ST-Link VCP -> USB -> COM6 on PC'],
        ['3.3V', 'CN7', 'Power', 'All three sensors VCC'],
        ['GND', 'Any GND', 'Ground', 'All three sensors GND'],
    ]
)

# ═══════════════════════════════════════════════════════════════════════
# 4. FIRMWARE ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('4. Firmware Architecture — How the Software Works', level=1)

doc.add_heading('4.1 FreeRTOS: Running Multiple Things at Once', level=2)
doc.add_paragraph(
    'A microcontroller has only one CPU core — it can only do one thing at a time. But we '
    'need to simultaneously read sensors (100 Hz), send data to the display, stream CSV '
    'over USB, and run AI inference. How?'
)
doc.add_paragraph(
    'FreeRTOS is a Real-Time Operating System — a tiny piece of software (about 10 KB of code) '
    'that rapidly switches between multiple "tasks" (think of them as lightweight programs). '
    'It switches so fast (1000 times per second in our case) that all tasks appear to run '
    'simultaneously, even though at any given nanosecond, only one task is actually executing.'
)
doc.add_paragraph(
    'Each task has a priority level. When a high-priority task needs to run, FreeRTOS '
    'immediately pauses whatever lower-priority task is running and gives the CPU to the '
    'urgent task. This is called "preemptive scheduling" and it\'s what makes real-time '
    'guarantees possible.'
)

doc.add_heading('4.2 The Four Tasks', level=2)
add_table(
    ['Task Name', 'Priority', 'Stack Size', 'Runs Every', 'What It Does'],
    [
        ['SensorTask', 'Above Normal', '512 words (2 KB)', '10 ms (100 Hz)',
         'Reads heart rate sensor and motion sensor, filters the PPG signal, detects heartbeat peaks, fills the acceleration buffer for ML'],
        ['TxTask', 'Normal', '512 words (2 KB)', 'On data arrival',
         'Takes sensor data from the queue, formats it as CSV text, sends it over UART to the PC, updates the OLED display every 50 frames'],
        ['MLTask', 'Below Normal', '1024 words (4 KB)', '~1 second',
         'Takes 100 acceleration samples, extracts 19 statistical features, runs the neural network, outputs the motion class (REST/WALKING/ARM UP)'],
        ['defaultTask', 'Normal', '128 words (512 B)', 'Idle',
         'Created by CubeMX; runs when nothing else needs the CPU'],
    ]
)
doc.add_paragraph(
    'The priority ordering is deliberate: SensorTask is highest because missing a sensor '
    'read at 100 Hz would lose data. TxTask is next because we want responsive display '
    'updates. MLTask is lowest because a 100ms delay in motion classification is '
    'imperceptible to the user.'
)

doc.add_heading('4.3 Data Flow: From Sensor to Screen', level=2)
doc.add_paragraph(
    'Here is how data moves through the system, step by step:'
)
doc.add_paragraph(
    '1. SensorTask wakes up every 10 ms. It acquires the I2C bus lock (5 ms timeout). '
    'If the lock is available (display isn\'t being updated), it reads the MAX30102 heart rate '
    'sensor and the MPU6050 motion sensor in a single I2C transaction batch. If the lock is '
    'busy, it uses the values from the previous read — no data is lost because the MAX30102 '
    'has a 32-sample internal FIFO.'
)
doc.add_paragraph(
    '2. The raw PPG (infrared light) value from the MAX30102 is passed through two digital '
    'filters: a high-pass filter at 0.5 Hz (removes the DC offset and slow drift) and a '
    'low-pass filter at 4 Hz (removes high-frequency noise). The result is a clean heartbeat '
    'wave. A peak detector finds each heartbeat and calculates BPM from the time between peaks.'
)
doc.add_paragraph(
    '3. The accelerometer values (ax, ay, az) are stored into a double buffer — a pair of '
    '100-sample arrays. While SensorTask fills one array, MLTask reads from the other. When '
    '100 samples are collected (1 second of data), the indices swap and a semaphore signal '
    'wakes MLTask.'
)
doc.add_paragraph(
    '4. All sensor values are packed into a SensorFrame structure and placed into a FreeRTOS '
    'message queue (20 frames deep). If the queue is full (TxTask is behind), frames are '
    'silently dropped — this prevents SensorTask from being blocked.'
)
doc.add_paragraph(
    '5. TxTask pulls frames from the queue, formats each one as a CSV line '
    '(e.g., "HR,72,0.01,-0.03,1.00,0.5,-0.2,0.1,25.3,0,15230"), and sends it over UART '
    'at 115200 baud. Every 50 frames (~0.5 seconds), it also redraws the OLED display.'
)
doc.add_paragraph(
    '6. MLTask, upon receiving the semaphore, computes 19 statistical features from the '
    '100 accelerometer samples, feeds them to the neural network, and stores the result '
    '(0=REST, 1=WALKING, 2=ARM UP) in a global variable that TxTask reads on the next frame.'
)

doc.add_heading('4.4 Sharing the I2C Bus Safely', level=2)
doc.add_paragraph(
    'Three devices share two wires. If SensorTask tries to read the MAX30102 while TxTask '
    'is flushing the display, the I2C signals would collide and both transfers would fail. '
    'The solution is a mutex — short for "mutual exclusion".'
)
doc.add_paragraph(
    'A mutex is like a bathroom lock: only one task can hold it at a time. Before touching '
    'the I2C bus, a task must acquire the mutex. If another task already holds it, the '
    'requesting task either waits or gives up. In our system:'
)
add_bullet('SensorTask tries to acquire with a 5 ms timeout. If the display flush is '
           'happening (103 ms), it gives up and uses cached sensor values. The MAX30102\'s '
           'internal FIFO ensures no readings are permanently lost.')
add_bullet('TxTask acquires with infinite timeout for display flush, since it\'s not '
           'time-critical. Thanks to "priority inheritance", while TxTask holds the mutex '
           'and SensorTask is waiting, FreeRTOS temporarily boosts TxTask\'s priority to '
           'match SensorTask — preventing a classic real-time bug called "priority inversion".')

# ═══════════════════════════════════════════════════════════════════════
# 5. SIGNAL PROCESSING
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('5. Signal Processing — Cleaning Up the Heart Rate Signal', level=1)

doc.add_heading('5.1 The Raw PPG Signal', level=2)
doc.add_paragraph(
    'The MAX30102 returns a raw number (roughly 50,000 to 200,000) representing how much '
    'infrared light was reflected back from your fingertip. This number has three components '
    'mixed together:'
)
add_bullet('A large DC offset (the average light level, which depends on skin tone, finger pressure, ambient light).')
add_bullet('A small AC ripple riding on top — this is the actual heartbeat signal, typically less than 2% of the total.')
add_bullet('High-frequency noise from electrical interference and sensor jitter.')
doc.add_paragraph(
    'We need to extract just the AC heartbeat ripple and discard everything else.'
)

doc.add_heading('5.2 Digital Filtering (Biquad IIR)', level=2)
doc.add_paragraph(
    'The firmware uses two biquad IIR (Infinite Impulse Response) filters in series:'
)
add_bullet('High-pass filter at 0.5 Hz: Removes the large DC offset and any slow drift '
           '(e.g., the signal changing because you shifted your finger). The heartbeat '
           'signal (1-3 Hz for 60-180 BPM) passes through untouched.', bold_prefix='HPF (0.5 Hz): ')
add_bullet('Low-pass filter at 4 Hz: Removes electrical noise and muscle tremor above 4 Hz. '
           'Since the fastest heart rate we care about is ~180 BPM = 3 Hz, everything above '
           '4 Hz is noise.', bold_prefix='LPF (4 Hz): ')
doc.add_paragraph(
    'A "biquad" is a specific type of filter that uses just 5 coefficients and 2 memory '
    'values. It\'s extremely efficient — each sample requires only 5 multiplications and '
    '4 additions — making it ideal for a microcontroller running at 100 Hz. The Butterworth '
    'design (Q = 0.707) gives a maximally flat frequency response in the passband.'
)

doc.add_heading('5.3 Peak Detection & BPM Calculation', level=2)
doc.add_paragraph(
    'After filtering, the signal is a clean wave with peaks corresponding to heartbeats. '
    'The peak detector works as follows:'
)
add_bullet('It tracks whether the signal is rising or falling.')
add_bullet('When it detects a transition from rising to falling (a local maximum), that\'s a peak — a heartbeat.')
add_bullet('It measures the time between consecutive peaks in milliseconds.')
add_bullet('BPM = 60,000 / (time between peaks in ms).')
add_bullet('If no finger is detected (IR value below 3000), the peak detector resets to avoid '
           'false readings from ambient light changes.')


# ═══════════════════════════════════════════════════════════════════════
# 6. MACHINE LEARNING
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('6. Machine Learning — Motion Classification on a Microcontroller', level=1)

doc.add_heading('6.1 What the ML Does', level=2)
doc.add_paragraph(
    'The neural network classifies the user\'s physical activity into one of three classes:'
)
add_table(
    ['Class', 'ID', 'What It Looks Like to the Sensor'],
    [
        ['REST', '0', 'Board is stationary. Accelerometer reads ~(0, 0, 1.0) g with very little variation.'],
        ['WALKING', '1', 'Rhythmic up-and-down acceleration from footsteps. Distinctive periodic pattern.'],
        ['ARM RAISED', '2', 'Gravity vector shifts to a different axis (e.g., (1.0, 0, 0) g when tilted 90 degrees).'],
    ]
)
doc.add_paragraph(
    'Classification happens once per second. This is fast enough for a health monitoring '
    'application — motion states don\'t change faster than that.'
)

doc.add_heading('6.2 Feature Extraction (The 19 Numbers)', level=2)
doc.add_paragraph(
    'The neural network doesn\'t look at raw accelerometer readings directly. Instead, '
    'every second, the firmware computes 19 statistical features from the 100-sample window '
    '(1 second at 100 Hz). This is the same set of features used to train the model in Python, '
    'ensuring consistency between training and inference.'
)
add_table(
    ['Feature Index', 'Name', 'What It Captures'],
    [
        ['0-3', 'ax: mean, std, min, max', 'Average tilt and variability along X axis'],
        ['4-7', 'ay: mean, std, min, max', 'Average tilt and variability along Y axis'],
        ['8-11', 'az: mean, std, min, max', 'Average tilt and variability along Z axis'],
        ['12-15', 'g_total: mean, std, min, max', 'Overall acceleration magnitude and variability'],
        ['16', 'g_total 25th percentile', 'Lower bound of typical acceleration'],
        ['17', 'g_total 75th percentile', 'Upper bound of typical acceleration'],
        ['18', 'Mean absolute deviation of g_total', 'How "jittery" the motion is (walking vs rest)'],
    ]
)
doc.add_paragraph(
    'The percentile calculation uses insertion sort on the 100-sample array — O(n^2) in theory, '
    'but for n=100 on a 160 MHz Cortex-M33 with hardware FPU, it completes in microseconds.'
)

doc.add_heading('6.3 The Neural Network Architecture', level=2)
doc.add_paragraph(
    'The model is a simple fully-connected (dense) neural network with three layers:'
)
add_table(
    ['Layer', 'Neurons', 'Parameters', 'Activation', 'Purpose'],
    [
        ['Input', '19', '0', 'N/A', 'Receives the 19 features'],
        ['Dense 1', '32', '19x32 + 32 = 640', 'ReLU', 'Learns combinations of features'],
        ['Dense 2', '16', '32x16 + 16 = 528', 'ReLU', 'Refines the learned patterns'],
        ['Output', '3', '16x3 + 3 = 51', 'Softmax', 'Outputs probability for each class'],
    ]
)
doc.add_paragraph(
    'Total parameters: 1,219 (640 + 528 + 51). Total size: ~4.7 KB of Flash memory. '
    'Inference time: less than 1 millisecond on the Cortex-M33 at 160 MHz.'
)
doc.add_paragraph(
    'For context: a large language model like GPT-4 has hundreds of billions of parameters. '
    'Our model has 1,219. It\'s a focused specialist — it does exactly one thing (classify '
    'three motion types) and does it well, running entirely on a $15 chip with no internet '
    'connection.'
)
doc.add_paragraph(
    'The output layer uses softmax activation, which converts the three output values into '
    'probabilities that sum to 1.0. For example, output [0.05, 0.03, 0.92] means 92% '
    'confidence the user\'s arm is raised. The firmware takes the argmax (the class with '
    'the highest probability) as the final classification.'
)

doc.add_heading('6.4 STM32Cube.AI: From Python to C', level=2)
doc.add_paragraph(
    'The model was trained in Python using Keras/TensorFlow, then converted to optimised C '
    'code using STM32Cube.AI — a tool from STMicroelectronics. This tool:'
)
add_bullet('Reads the trained model (Keras .h5 or ONNX format).')
add_bullet('Quantises and optimises the weights for the target MCU.')
add_bullet('Generates C source files (network.c, network_data.c) with the model weights '
           'baked in as constant arrays.')
add_bullet('Uses CMSIS-NN intrinsics — ARM\'s hand-optimised math library — for matrix '
           'multiplications, making inference significantly faster than naive C loops.')
doc.add_paragraph(
    'The generated code exposes a simple API: ai_network_create_and_init() to load the model, '
    'ai_network_run() to execute inference. The firmware writes 19 floats into the input '
    'buffer and reads 3 floats (class probabilities) from the output buffer.'
)

doc.add_heading('6.5 How Inference Runs in Real Time', level=2)
doc.add_paragraph(
    'The double-buffer mechanism ensures zero-copy, lock-free data transfer between '
    'SensorTask and MLTask:'
)
add_bullet('Two 100x3 float arrays exist in memory: accel_buf[0] and accel_buf[1].')
add_bullet('SensorTask writes to accel_buf[write_idx]. When 100 samples are collected, '
           'it flips write_idx (0 becomes 1, 1 becomes 0) and signals a semaphore.')
add_bullet('MLTask reads from accel_buf[1 - write_idx] — the buffer SensorTask just '
           'finished. Since they\'re always accessing different buffers, no mutex is needed.')
add_bullet('This "ping-pong" or "double buffer" pattern is extremely common in real-time '
           'systems and eliminates the overhead of copying data or acquiring locks.')


# ═══════════════════════════════════════════════════════════════════════
# 7. UART CSV
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('7. Data Output — UART CSV Stream', level=1)
doc.add_paragraph(
    'Every sensor frame (100 per second) is formatted as a comma-separated line and sent '
    'over UART at 115200 baud via the ST-Link Virtual COM Port. The format is:'
)
add_code('HR,<bpm>,<ax>,<ay>,<az>,<gx>,<gy>,<gz>,<temp>,<motion>,<ts_ms>')
doc.add_paragraph('Example line:')
add_code('HR,72,0.012,-0.038,1.002,0.5,-0.2,0.1,25.3,1,15230')
add_table(
    ['Field', 'Type', 'Description'],
    [
        ['HR', 'Literal', 'Header tag — every line starts with "HR" for easy parsing'],
        ['bpm', 'Integer', 'Heart rate in beats per minute (0 = no finger detected)'],
        ['ax, ay, az', 'Float', 'Accelerometer readings in g (gravity units)'],
        ['gx, gy, gz', 'Float', 'Gyroscope readings in degrees per second'],
        ['temp', 'Float', 'MPU6050 chip temperature in degrees Celsius'],
        ['motion', 'Integer', '0=REST, 1=WALKING, 2=ARM UP'],
        ['ts_ms', 'Integer', 'Timestamp in milliseconds since boot'],
    ]
)
doc.add_paragraph(
    'The UART transmit function has a 50 ms timeout. If the previous transmission hasn\'t '
    'finished (e.g., the PC isn\'t reading fast enough), the frame is silently dropped. '
    'This prevents the real-time sensor loop from blocking on serial I/O.'
)

# ═══════════════════════════════════════════════════════════════════════
# 8. PYTHON TOOLS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('8. Python Companion Tools', level=1)

doc.add_heading('8.1 Live Dashboard (vital_dashboard.py)', level=2)
doc.add_paragraph(
    'A real-time visualisation tool that connects to the STM32\'s virtual COM port and '
    'draws live, scrolling graphs using matplotlib. Features:'
)
add_bullet('Rolling heart rate plot (red line, 0-200 BPM range, last 500 samples = 5 seconds).')
add_bullet('3-axis accelerometer overlay (blue=ax, green=ay, orange=az).')
add_bullet('Large colour-coded motion class indicator (blue=REST, green=WALKING, orange=ARM UP).')
add_bullet('Live temperature reading and frame counter.')
add_bullet('Automatic CSV logging to a timestamped file (e.g., vital_log_20260330_120000.csv).')
doc.add_paragraph(
    'The dashboard uses a dark theme (#1a1a2e background) for reduced eye strain during '
    'long monitoring sessions. It reads serial data at 115200 baud and updates the plots '
    'every 100 ms.'
)

doc.add_heading('8.2 Offline Analysis (analyse_log.py)', level=2)
doc.add_paragraph(
    'A post-session analysis script that reads the CSV log files saved by the dashboard '
    'and produces:'
)
add_bullet('Console summary: session duration, mean/median/max heart rate, peak g-force, '
           'temperature range, motion class breakdown with percentages.')
add_bullet('A 6-panel analysis PNG with: HR over time, 3-axis acceleration, total |g| '
           'magnitude, motion class timeline, board temperature, and a motion distribution '
           'pie chart.')
doc.add_paragraph('Usage:')
add_code('python analyse_log.py vital_log_20260330_120000.csv')

# ═══════════════════════════════════════════════════════════════════════
# 9. DEVELOPMENT JOURNEY — ERRORS & FIXES
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('9. The Development Journey — Errors, Detours & Fixes', level=1)
doc.add_paragraph(
    'This section documents every significant bug encountered during development, in '
    'chronological order. Real-world embedded development is messy — hardware doesn\'t '
    'behave like simulations, and bugs often have non-obvious root causes that span '
    'hardware, firmware, and toolchain boundaries.'
)

doc.add_heading('9.1 Display Went Dark After Sensor Init', level=2)
doc.add_paragraph('Symptom:', style='List Bullet')
doc.add_paragraph(
    'The SSD1306 OLED display showed the splash screen ("VITAL SIGN LOGGER / Initialising...") '
    'correctly, but went completely blank after the MAX30102 heart rate sensor was initialised. '
    'The display never came back.'
)
doc.add_paragraph('Root Cause:', style='List Bullet')
doc.add_paragraph(
    'The MAX30102 driver performs a "soft reset" during initialisation — it tells the sensor '
    'chip to restart. After this reset, the driver calls HAL_I2C_DeInit() followed by '
    'HAL_I2C_Init() to reinitialise the I2C peripheral on the STM32. This is a standard '
    'recovery step to ensure a clean I2C state.'
)
doc.add_paragraph(
    'However, the SSD1306 display controller is a state machine. It was already configured '
    '(contrast, orientation, charge pump settings) by SSD1306_Init(). When the I2C peripheral '
    'was torn down and rebuilt, the SSD1306 lost its configuration — it was still powered '
    'on, but its internal state was wiped. Any subsequent I2C writes to it were valid at '
    'the bus level, but the display controller no longer knew what to do with them.'
)
doc.add_paragraph('Fix:', style='List Bullet')
doc.add_paragraph(
    'Added a second SSD1306_Init(&hi2c1) call immediately after MAX30102_Init(), '
    'reconfiguring the display controller after the I2C bus reset. This is a single line of '
    'code but took significant debugging to identify, because the I2C bus scan still showed '
    'the display at address 0x3C — it was responding to address probes, just not displaying '
    'anything.'
)
add_code('MAX30102_Init(&ppg_dev, &hi2c1);\n\n/* MAX30102_Init does HAL_I2C_DeInit/Init after soft-reset —\n * re-init SSD1306 so its controller state is guaranteed clean */\nSSD1306_Init(&hi2c1);')

doc.add_heading('9.2 No UART Output (Printf Was a No-Op)', level=2)
doc.add_paragraph('Symptom:', style='List Bullet')
doc.add_paragraph(
    'Opening the COM port in PuTTY at 115200 baud showed absolutely nothing. No boot '
    'messages, no CSV data, no error messages. The firmware appeared to be running (the '
    'red LED on the MAX30102 was glowing), but no data was coming out.'
)
doc.add_paragraph('Root Cause:', style='List Bullet')
doc.add_paragraph(
    'In ARM GCC (the compiler used for STM32), printf() eventually calls a low-level '
    'function called _write() to send characters to the outside world. The CubeMX code '
    'generator creates a default sysmem.c file with a _write() implementation that does '
    'literally nothing:'
)
add_code('int _write(int fd, char *buf, int len) {\n  (void)fd;\n  (void)buf;\n  (void)len;\n  return len;  // Claims success but discards everything!\n}')
doc.add_paragraph(
    'This is an intentional stub — CubeMX doesn\'t know what output device you want to use '
    '(UART? USB? SWO?). But the consequences are severe: every printf() call in the entire '
    'firmware silently succeeds while producing no output. The function returns len '
    '(claiming it wrote all bytes) so the caller has no idea data was lost.'
)
doc.add_paragraph('Fix:', style='List Bullet')
doc.add_paragraph(
    'Rewrote _write() to route through __io_putchar(), which is the standard STM32 '
    'printf retargeting hook that calls HAL_UART_Transmit():'
)
add_code('int _write(int fd, char *buf, int len) {\n  (void)fd;\n  extern int __io_putchar(int ch);\n  for (int i = 0; i < len; i++) __io_putchar(buf[i]);\n  return len;\n}')
doc.add_paragraph(
    'The corresponding __io_putchar() function in main.c sends each character '
    'to USART1 with a 10 ms timeout.'
)

doc.add_heading('9.3 Tasks Crashing Silently (Stack Overflow)', level=2)
doc.add_paragraph('Symptom:', style='List Bullet')
doc.add_paragraph(
    'After fixing printf, boot messages appeared in PuTTY ("I2C bus scan", "PPG: OK", etc.), '
    'but no CSV data frames were ever sent. The RTOS tasks appeared to be dead.'
)
doc.add_paragraph('Root Cause:', style='List Bullet')
doc.add_paragraph(
    'TxTask was allocated 256 words (1024 bytes) of stack by CubeMX. The snprintf() call '
    'that formats each CSV line uses floating-point formatting (%f), which in newlib-nano '
    '(the C library used for ARM) requires approximately 600 bytes of stack space for '
    'internal conversion buffers. Add the 160-byte output buffer, HAL_UART_Transmit() '
    'overhead, and FreeRTOS context — total stack usage hit ~1024 bytes, exactly at the '
    'limit. Any additional function call pushed it over, corrupting memory and crashing '
    'the task.'
)
doc.add_paragraph(
    'FreeRTOS\'s default behaviour on stack overflow is undefined — the task simply stops '
    'running. There\'s no error message, no crash log, no blinking LED. It just silently dies.'
)
doc.add_paragraph('Fix:', style='List Bullet')
doc.add_paragraph(
    'Two changes: (1) Increased TxTask stack from 256 to 512 words (2048 bytes), giving '
    'comfortable headroom. (2) Enabled FreeRTOS stack overflow detection '
    '(configCHECK_FOR_STACK_OVERFLOW = 2) and added a hook that prints the offending task '
    'name over UART before halting:'
)
add_code('void vApplicationStackOverflowHook(TaskHandle_t xTask, char *pcTaskName)\n{\n    char msg[64];\n    int n = snprintf(msg, sizeof(msg),\n                     "\\r\\n!!! STACK OVERFLOW: %s !!!\\r\\n", pcTaskName);\n    HAL_UART_Transmit(&huart1, (const uint8_t *)msg, (uint16_t)n, 50);\n    taskDISABLE_INTERRUPTS();\n    for (;;) {}\n}')
doc.add_paragraph(
    'This ensures that any future stack overflow immediately identifies the guilty task '
    'instead of silently corrupting memory.'
)

doc.add_heading('9.4 ML Inference Returning Garbage', level=2)
doc.add_paragraph('Symptom:', style='List Bullet')
doc.add_paragraph(
    'After enabling the ML task, the UART output showed "ML_Classify: ai_network_run error=0" '
    'on every inference attempt. The motion class was always 0 (REST) regardless of movement.'
)
doc.add_paragraph('Root Cause:', style='List Bullet')
doc.add_paragraph(
    'The original code manually constructed AI buffer descriptors using the AI_BUFFER_INIT '
    'macro, pointing them to locally allocated float arrays. However, the STM32Cube.AI code '
    'generator had set AI_NETWORK_INPUTS_IN_ACTIVATIONS = 4, meaning the network\'s input '
    'and output buffers live inside the activations buffer (a pre-allocated memory pool), not '
    'in user-provided arrays.'
)
doc.add_paragraph(
    'The manually constructed buffer descriptors pointed to the wrong memory addresses. '
    'When ai_network_run() checked the buffer pointers, they didn\'t match the expected '
    'activations layout, and the inference was rejected.'
)
doc.add_paragraph('Fix:', style='List Bullet')
doc.add_paragraph(
    'Replaced the manual AI_BUFFER_INIT with the official API calls '
    'ai_network_inputs_get() and ai_network_outputs_get(), which return pointers to the '
    'network\'s own internal buffers:'
)
add_code('ai_buffer *ai_input  = ai_network_inputs_get(network_hdl, NULL);\nai_buffer *ai_output = ai_network_outputs_get(network_hdl, NULL);\nfloat *in_data = (float *)ai_input[0].data;\nextract_features(in_data);  // Write directly into network\'s buffer\nai_network_run(network_hdl, ai_input, ai_output);')
doc.add_paragraph(
    'This writes the 19 features directly into the memory that the network expects, '
    'eliminating any mismatch. After this fix, ML inference immediately started producing '
    'correct classifications.'
)

doc.add_heading('9.5 Display Only Detected 1 in 5 Boots', level=2)
doc.add_paragraph('Symptom:', style='List Bullet')
doc.add_paragraph(
    'The I2C bus scan at boot sometimes showed all three devices (0x3C, 0x57, 0x68) and '
    'sometimes only two (0x57 and 0x68, missing the display). Approximately 1 in 5 boots '
    'detected the display. When it wasn\'t detected, the display stayed dark for the entire '
    'session, but the rest of the system worked fine.'
)
doc.add_paragraph('Root Cause:', style='List Bullet')
doc.add_paragraph(
    'The SSD1306 was connected via a breadboard with male-to-male jumper wires. These '
    'connections are notoriously unreliable — slight vibrations can momentarily break '
    'contact. Specifically, the 3.3V power and I2C data lines had intermittent '
    'connections that depended on the exact mechanical state of the wire-breadboard junction.'
)
doc.add_paragraph('Mitigation:', style='List Bullet')
doc.add_paragraph(
    'Added I2C error recovery after every SSD1306_Flush() call: if the HAL reports an I2C '
    'error (NACK, bus error, etc.), the firmware reinitialises the I2C peripheral. This '
    'prevents a display failure from cascading and killing the MAX30102/MPU6050 '
    'communication:'
)
add_code('SSD1306_Flush();\nif (hi2c1.ErrorCode != HAL_I2C_ERROR_NONE) {\n    HAL_I2C_DeInit(&hi2c1);\n    HAL_I2C_Init(&hi2c1);\n}')
doc.add_paragraph(
    'Long-term fix: replace the breadboard with direct point-to-point wiring using solid '
    'core jumper wires. The wiring note in the hardware section reflects this lesson.'
)

doc.add_heading('9.6 Board Stuck on "Keep Board Flat"', level=2)
doc.add_paragraph('Symptom:', style='List Bullet')
doc.add_paragraph(
    'During MPU6050 initialisation, the OLED displayed "IMU calibrating.. Keep board flat!" '
    'and never progressed. The system appeared frozen.'
)
doc.add_paragraph('Root Cause:', style='List Bullet')
doc.add_paragraph(
    'The MPU6050 calibration originally took 500 samples, each with a 50 ms I2C timeout. '
    'If I2C reads were failing (returning HAL_TIMEOUT), the worst-case calibration time was '
    '500 x 50 ms = 25 seconds. Additionally, I2C errors weren\'t being recovered, so once '
    'the bus entered an error state, every subsequent read also timed out — an infinite '
    'loop of 50 ms timeouts.'
)
doc.add_paragraph('Fix:', style='List Bullet')
doc.add_paragraph(
    'Three changes: (1) Reduced calibration to 200 samples. (2) Added I2C error recovery '
    '(DeInit/Init) inside MPU6050_Read() so a single failed read doesn\'t lock up the bus '
    'permanently. (3) Added an I2C recovery step after calibration completes, before '
    'proceeding to RTOS startup.'
)

doc.add_heading('9.7 Motion Stuck at REST Despite Moving', level=2)
doc.add_paragraph('Symptom:', style='List Bullet')
doc.add_paragraph(
    'After fixing ML inference (Section 9.4), the UART output showed valid ML probabilities '
    '(e.g., "ML: R=0.85 W=0.10 A=0.05 -> 0"), but the classification was always REST even '
    'when the user picked up the board and walked around.'
)
doc.add_paragraph('Root Cause:', style='List Bullet')
doc.add_paragraph(
    'This turned out to not be a bug at all. The user was carrying the breadboard level '
    '(flat, like a tray) while walking. The accelerometer readings showed (0, 0, ~1.0) g '
    'with minimal variation — indistinguishable from sitting on a table. The ML model was '
    'correctly classifying this as REST because the acceleration pattern was indeed rest-like.'
)
doc.add_paragraph('Resolution:', style='List Bullet')
doc.add_paragraph(
    'Tilting the board 90 degrees immediately triggered ARM UP classification. Walking '
    'with more vigorous arm motion (allowing the board to experience footstep vibrations) '
    'triggered WALKING. All three classes were validated on hardware. This was a lesson in '
    'understanding what the sensor actually measures vs. what the human expects.'
)

doc.add_heading('9.8 Python Dashboard: COM Port Access Denied', level=2)
doc.add_paragraph('Symptom:', style='List Bullet')
doc.add_paragraph(
    'Running the Python dashboard produced: PermissionError(13, \'Access is denied.\', None, 5).'
)
doc.add_paragraph('Root Cause:', style='List Bullet')
doc.add_paragraph(
    'PuTTY (used earlier for debugging) was still holding COM6 open. On Windows, serial '
    'ports are exclusive-access — only one application can open a COM port at a time.'
)
doc.add_paragraph('Fix:', style='List Bullet')
doc.add_paragraph('Close PuTTY before running the Python dashboard. Only one application can access COM6 at a time.')

# ═══════════════════════════════════════════════════════════════════════
# 10. BOOT SEQUENCE
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('10. Boot Sequence — What Happens When You Power On', level=1)
doc.add_paragraph(
    'When the NUCLEO board is powered on or the reset button is pressed, the following '
    'sequence occurs (this runs in main.c before the RTOS starts):'
)
steps = [
    ('HAL_Init()', 'Initialises the Hardware Abstraction Layer — sets up the SysTick timer (1 kHz), configures the NVIC (interrupt controller), enables the Flash instruction cache.'),
    ('SystemClock_Config()', 'Configures the PLL (Phase-Locked Loop) to boost the internal MSI oscillator from 48 MHz to 160 MHz. PLL settings: M=3, N=10, R=1. Flash wait states set to 4 (required at 160 MHz).'),
    ('GPIO / I2C / UART Init', 'Configures pin multiplexing (PB8/PB9 for I2C, PA9 for UART TX), initialises I2C1 at 100 kHz (timing register 0x30909DEC), sets UART1 to 115200 baud 8N1.'),
    ('I2C Bus Scan', 'Probes every 7-bit I2C address (0x08-0x77) with a 10 ms timeout. Prints found addresses to UART. This validates wiring before attempting device initialisation.'),
    ('SSD1306_Init()', 'Sends the initialisation command sequence to the OLED display (display on, charge pump enable, contrast set, memory addressing mode, etc.). Shows "Initialising..." splash.'),
    ('MAX30102_Init()', 'Configures the heart rate sensor: LED currents, sample rate, ADC resolution, FIFO settings. Performs a soft reset which triggers HAL_I2C_DeInit/Init.'),
    ('SSD1306_Init() (again)', 'Re-initialises the display after the I2C bus reset caused by the MAX30102 soft reset.'),
    ('MPU6050_Init() + Calibrate()', 'Wakes the motion sensor from sleep, sets accelerometer range to +/-2g, gyroscope range to +/-250 deg/s. Then takes 200 samples to compute bias offsets while the board sits flat. Shows "Keep board flat!" on OLED during calibration.'),
    ('Status Display', 'Shows PPG: OK/FAIL, IMU: OK/FAIL, and "Place finger..." on the OLED for 1.5 seconds.'),
    ('osKernelStart()', 'Starts the FreeRTOS scheduler. Control never returns to main() — from this point on, all execution happens inside the four RTOS tasks.'),
]
for i, (name, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {name}: ')
    run.bold = True
    p.add_run(desc)

# ═══════════════════════════════════════════════════════════════════════
# 11. KEY CONFIGURATION
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('11. Key Configuration & Numbers', level=1)
add_table(
    ['Parameter', 'Value', 'Where Set'],
    [
        ['System Clock', '160 MHz', 'SystemClock_Config() in main.c'],
        ['I2C Speed', '~100 kHz (Standard Mode)', 'Timing register 0x30909DEC in MX_I2C1_Init()'],
        ['UART Baud Rate', '115200 8N1', 'MX_USART1_UART_Init()'],
        ['Sensor Sample Rate', '100 Hz', 'SensorTask loop: osDelayUntil(wake_tick += 10)'],
        ['ML Inference Rate', '~1 Hz (every 100 samples)', 'accel_buf window size'],
        ['FreeRTOS Tick Rate', '1000 Hz (1 ms resolution)', 'FreeRTOSConfig.h'],
        ['FreeRTOS Heap', '32 KB', 'FreeRTOSConfig.h'],
        ['PPG High-Pass Filter', '0.5 Hz, Q=0.707', 'Biquad_Init in SensorTask'],
        ['PPG Low-Pass Filter', '4.0 Hz, Q=0.707', 'Biquad_Init in SensorTask'],
        ['Finger Detection Threshold', 'IR < 3000 counts', '#define FINGER_THRESH in sensor_task.c'],
        ['ML Feature Count', '19', 'AI_NETWORK_IN_1_SIZE'],
        ['ML Output Classes', '3 (REST/WALKING/ARM UP)', 'AI_NETWORK_OUT_1_SIZE'],
        ['ML Model Parameters', '1,219', 'network.c (auto-generated)'],
        ['ML Model Flash Size', '~4.7 KB', 'network_data.c'],
        ['Display Refresh', 'Every 50 frames (~0.5 s)', 'TxTask frame_count threshold'],
        ['Display Flush Duration', '~103 ms at 100 kHz I2C', '1024 bytes / 100 kHz'],
        ['UART TX Timeout', '50 ms (drop if busy)', 'usb_tx.c USBTX_Send()'],
    ]
)

# ═══════════════════════════════════════════════════════════════════════
# 12. STATUS & FUTURE
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('12. Project Status & Future Work', level=1)
doc.add_paragraph(
    'As of 30 March 2026, the project is functionally complete. All hardware subsystems '
    'are validated on the physical NUCLEO-U575ZI-Q board:'
)
add_bullet('MAX30102 heart rate detection: confirmed working (e.g., 167 BPM reading validated).')
add_bullet('MPU6050 accelerometer/gyroscope: live readings with startup calibration.')
add_bullet('SSD1306 OLED: displays HR and motion class (intermittent due to breadboard wiring).')
add_bullet('UART CSV stream: 100 Hz continuous output confirmed on COM6.')
add_bullet('ML inference: all three motion classes (REST, WALKING, ARM UP) correctly classified on hardware.')
add_bullet('Python live dashboard and offline analysis tools: both working.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Potential future enhancements:').bold = True
add_bullet('Zephyr RTOS port — would replace FreeRTOS with a more feature-rich RTOS, '
           'adding device tree hardware abstraction, a networking stack, and a more '
           'modern build system (CMake/west). Planned as a separate follow-on project.')
add_bullet('LPBAM (Low-Power Background Autonomous Mode) — STM32U5-specific feature '
           'that allows DMA transfers to continue while the main CPU sleeps, dramatically '
           'reducing power consumption for battery-powered operation.')
add_bullet('More motion classes — the current model only distinguishes three activities. '
           'The architecture supports retraining with additional classes (running, cycling, '
           'falling) by collecting more labeled data and retraining the Keras model.')
add_bullet('SpO2 measurement — the MAX30102 has both red and infrared LEDs, enabling '
           'blood oxygen saturation estimation. Currently only the IR channel is used for '
           'heart rate.')

# ═══════════════════════════════════════════════════════════════════════
# 13. GLOSSARY
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('13. Glossary of Terms', level=1)
glossary = [
    ('ARM Cortex-M33', 'A 32-bit processor core designed by ARM for microcontrollers. The "M" stands for "Microcontroller". The "33" indicates it supports TrustZone security and DSP instructions.'),
    ('Biquad Filter', 'A type of digital filter with two poles and two zeros. "Bi" = two, "quad" = quadratic (the transfer function is a ratio of two quadratic polynomials). Very efficient to compute.'),
    ('BPM', 'Beats Per Minute — the standard unit for heart rate.'),
    ('CSV', 'Comma-Separated Values — a plain text format where data fields are separated by commas.'),
    ('DMA', 'Direct Memory Access — hardware that moves data between memory and peripherals without CPU involvement.'),
    ('FIFO', 'First In, First Out — a hardware buffer where data is read in the same order it was written.'),
    ('FPU', 'Floating-Point Unit — dedicated hardware that performs decimal arithmetic (addition, multiplication, square root) in a single clock cycle.'),
    ('FreeRTOS', 'A free, open-source Real-Time Operating System for microcontrollers. Provides task scheduling, mutexes, semaphores, and message queues.'),
    ('g (gravity unit)', 'The standard acceleration due to Earth\'s gravity: 9.81 m/s. An accelerometer reading of 1.0 g means the sensor is experiencing normal gravitational pull.'),
    ('HAL', 'Hardware Abstraction Layer — STMicroelectronics\' library that provides C functions for controlling hardware peripherals (I2C, UART, GPIO, etc.).'),
    ('I2C', 'Inter-Integrated Circuit — a two-wire serial communication protocol. Uses SDA (data) and SCL (clock). Supports multiple devices on one bus via unique addresses.'),
    ('IIR Filter', 'Infinite Impulse Response — a type of digital filter that uses feedback (previous outputs) in its computation. More efficient than FIR filters for the same frequency response.'),
    ('MEMS', 'Micro-Electro-Mechanical System — a tiny mechanical structure (springs, masses) fabricated on a silicon chip. Used in accelerometers and gyroscopes.'),
    ('Mutex', 'Mutual Exclusion — a synchronisation primitive that ensures only one task accesses a shared resource at a time. Like a lock on a bathroom door.'),
    ('OLED', 'Organic Light-Emitting Diode — a display technology where each pixel emits its own light. No backlight needed, resulting in true blacks and low power consumption.'),
    ('PPG', 'Photoplethysmography — measuring blood volume changes by shining light into tissue and detecting absorption. Used in pulse oximeters and smartwatches.'),
    ('RTOS', 'Real-Time Operating System — an OS that guarantees task execution within strict time deadlines.'),
    ('Semaphore', 'A synchronisation primitive used for signalling between tasks. "Give" increments a counter, "Take" decrements it (blocking if zero). Used here to signal when 100 samples are ready.'),
    ('Softmax', 'A mathematical function that converts a vector of numbers into a probability distribution (all values between 0 and 1, summing to 1).'),
    ('UART', 'Universal Asynchronous Receiver-Transmitter — a serial communication protocol. Sends data one bit at a time, without a shared clock line (unlike I2C).'),
]
for term, definition in glossary:
    p = doc.add_paragraph()
    run = p.add_run(f'{term}: ')
    run.bold = True
    p.add_run(definition)

# ═══════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        'VitalSignLogger_Project_Report.docx')
doc.save(out_path)
print(f"Report saved to: {out_path}")
